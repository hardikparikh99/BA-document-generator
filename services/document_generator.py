"""
Document generator service for the CrewAI Multi-Agent Project Documentation System.
This module handles document generation in various formats (PDF, DOCX, HTML).
"""
import os
import asyncio
import tempfile
from typing import Dict, Any, Optional, List
from datetime import datetime
import base64

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# HTML generation
import jinja2

from utils.config import get_settings, get_temp_dir
from utils.logger import setup_logger
from models.database import get_documentation_by_id, store_download_info

# Setup logger
logger = setup_logger(__name__)
settings = get_settings()

class DocumentGenerator:
    """
    Service for generating documents in various formats.
    """
    
    def __init__(self):
        """Initialize the document generator."""
        self.temp_dir = get_temp_dir()
        self.base_url = "http://localhost:7000"  # For download URLs
        
        # Set up Jinja2 for HTML templates
        template_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates")
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(template_dir),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
    
    async def generate_pdf(self, documentation: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a PDF document from the documentation.
        
        Args:
            documentation: Documentation data
            
        Returns:
            dict: Result of the operation with success status and file_path
        """
        try:
            # Create a unique filename
            filename = f"{documentation['file_id']}_documentation.pdf"
            file_path = os.path.join(self.temp_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Title'],
                fontSize=16,
                spaceAfter=12
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=10
            )
            
            normal_style = styles['Normal']
            
            # Content elements
            elements = []
            
            # Title
            elements.append(Paragraph(documentation['title'], title_style))
            elements.append(Spacer(1, 12))
            
            # Date
            date_str = datetime.now().strftime("%Y-%m-%d")
            elements.append(Paragraph(f"Generated on: {date_str}", normal_style))
            elements.append(Spacer(1, 24))
            
            # Executive Summary
            elements.append(Paragraph("Executive Summary", heading_style))
            elements.append(Paragraph(documentation['executive_summary'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Project Scope and Objectives
            elements.append(Paragraph("Project Scope and Objectives", heading_style))
            elements.append(Paragraph(documentation['project_scope'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Stakeholder Analysis
            elements.append(Paragraph("Stakeholder Analysis", heading_style))
            elements.append(Paragraph(documentation['stakeholder_analysis'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Page break
            elements.append(PageBreak())
            
            # Functional Requirements
            elements.append(Paragraph("Functional Requirements", heading_style))
            elements.append(Paragraph(documentation['functional_requirements'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Technical Requirements
            elements.append(Paragraph("Technical Requirements", heading_style))
            elements.append(Paragraph(documentation['technical_requirements'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Timeline and Milestones
            elements.append(Paragraph("Timeline and Milestones", heading_style))
            elements.append(Paragraph(documentation['timeline'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Page break
            elements.append(PageBreak())
            
            # Budget Considerations
            elements.append(Paragraph("Budget Considerations", heading_style))
            elements.append(Paragraph(documentation['budget'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Risk Assessment
            elements.append(Paragraph("Risk Assessment", heading_style))
            elements.append(Paragraph(documentation['risk_assessment'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Assumptions and Dependencies
            elements.append(Paragraph("Assumptions and Dependencies", heading_style))
            elements.append(Paragraph(documentation['assumptions'], normal_style))
            elements.append(Spacer(1, 12))
            
            # Next Steps and Recommendations
            elements.append(Paragraph("Next Steps and Recommendations", heading_style))
            elements.append(Paragraph(documentation['next_steps'], normal_style))
            
            # Build the PDF
            doc.build(elements)
            
            logger.info(f"PDF document generated successfully: {file_path}")
            
            # Create download URL
            download_url = f"{self.base_url}/static/downloads/{filename}"
            
            # Store download info
            await store_download_info({
                "file_id": documentation['file_id'],
                "documentation_id": documentation['documentation_id'],
                "format": "pdf",
                "file_path": file_path,
                "download_url": download_url,
                "expiry_time": datetime.now().isoformat()
            })
            
            return {
                "success": True,
                "file_path": file_path,
                "download_url": download_url
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF document: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating PDF document: {str(e)}"
            }
    
    async def generate_docx(self, documentation: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a DOCX document from the documentation.
        
        Args:
            documentation: Documentation data
            
        Returns:
            dict: Result of the operation with success status and file_path
        """
        try:
            # Create a unique filename
            filename = f"{documentation['file_id']}_documentation.docx"
            file_path = os.path.join(self.temp_dir, filename)
            
            # Create DOCX document
            doc = Document()
            
            # Title
            doc.add_heading(documentation['title'], level=0)
            
            # Date
            date_str = datetime.now().strftime("%Y-%m-%d")
            doc.add_paragraph(f"Generated on: {date_str}")
            doc.add_paragraph()
            
            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(documentation['executive_summary'])
            doc.add_paragraph()
            
            # Project Scope and Objectives
            doc.add_heading("Project Scope and Objectives", level=1)
            doc.add_paragraph(documentation['project_scope'])
            doc.add_paragraph()
            
            # Stakeholder Analysis
            doc.add_heading("Stakeholder Analysis", level=1)
            doc.add_paragraph(documentation['stakeholder_analysis'])
            doc.add_paragraph()
            
            # Page break
            doc.add_page_break()
            
            # Functional Requirements
            doc.add_heading("Functional Requirements", level=1)
            doc.add_paragraph(documentation['functional_requirements'])
            doc.add_paragraph()
            
            # Technical Requirements
            doc.add_heading("Technical Requirements", level=1)
            doc.add_paragraph(documentation['technical_requirements'])
            doc.add_paragraph()
            
            # Timeline and Milestones
            doc.add_heading("Timeline and Milestones", level=1)
            doc.add_paragraph(documentation['timeline'])
            doc.add_paragraph()
            
            # Page break
            doc.add_page_break()
            
            # Budget Considerations
            doc.add_heading("Budget Considerations", level=1)
            doc.add_paragraph(documentation['budget'])
            doc.add_paragraph()
            
            # Risk Assessment
            doc.add_heading("Risk Assessment", level=1)
            doc.add_paragraph(documentation['risk_assessment'])
            doc.add_paragraph()
            
            # Assumptions and Dependencies
            doc.add_heading("Assumptions and Dependencies", level=1)
            doc.add_paragraph(documentation['assumptions'])
            doc.add_paragraph()
            
            # Next Steps and Recommendations
            doc.add_heading("Next Steps and Recommendations", level=1)
            doc.add_paragraph(documentation['next_steps'])
            
            # Save the document
            doc.save(file_path)
            
            logger.info(f"DOCX document generated successfully: {file_path}")
            
            # Create download URL
            download_url = f"{self.base_url}/static/downloads/{filename}"
            
            # Store download info
            await store_download_info({
                "file_id": documentation['file_id'],
                "documentation_id": documentation['documentation_id'],
                "format": "docx",
                "file_path": file_path,
                "download_url": download_url,
                "expiry_time": datetime.now().isoformat()
            })
            
            return {
                "success": True,
                "file_path": file_path,
                "download_url": download_url
            }
            
        except Exception as e:
            logger.error(f"Error generating DOCX document: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating DOCX document: {str(e)}"
            }
    
    async def generate_html(self, documentation: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate an HTML document from the documentation.
        
        Args:
            documentation: Documentation data
            
        Returns:
            dict: Result of the operation with success status and file_path
        """
        try:
            # Create a unique filename
            filename = f"{documentation['file_id']}_documentation.html"
            file_path = os.path.join(self.temp_dir, filename)
            
            # Load HTML template
            try:
                template = self.jinja_env.get_template("document_template.html")
            except jinja2.exceptions.TemplateNotFound:
                # Create a basic template if not found
                template_str = """
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>{{ title }}</title>
                    <style>
                        body {
                            font-family: Arial, sans-serif;
                            line-height: 1.6;
                            margin: 0;
                            padding: 20px;
                            color: #333;
                        }
                        .container {
                            max-width: 800px;
                            margin: 0 auto;
                        }
                        h1 {
                            color: #2c3e50;
                            border-bottom: 2px solid #eee;
                            padding-bottom: 10px;
                        }
                        h2 {
                            color: #3498db;
                            margin-top: 30px;
                        }
                        .date {
                            color: #7f8c8d;
                            font-style: italic;
                            margin-bottom: 30px;
                        }
                        .section {
                            margin-bottom: 30px;
                        }
                        .footer {
                            margin-top: 50px;
                            padding-top: 20px;
                            border-top: 1px solid #eee;
                            font-size: 0.8em;
                            color: #7f8c8d;
                        }
                    </style>
                </head>
                <body>
                    <div class="container">
                        <h1>{{ title }}</h1>
                        <div class="date">Generated on: {{ date }}</div>
                        
                        <div class="section">
                            <h2>Executive Summary</h2>
                            <p>{{ executive_summary }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Project Scope and Objectives</h2>
                            <p>{{ project_scope }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Stakeholder Analysis</h2>
                            <p>{{ stakeholder_analysis }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Functional Requirements</h2>
                            <p>{{ functional_requirements }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Technical Requirements</h2>
                            <p>{{ technical_requirements }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Timeline and Milestones</h2>
                            <p>{{ timeline }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Budget Considerations</h2>
                            <p>{{ budget }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Risk Assessment</h2>
                            <p>{{ risk_assessment }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Assumptions and Dependencies</h2>
                            <p>{{ assumptions }}</p>
                        </div>
                        
                        <div class="section">
                            <h2>Next Steps and Recommendations</h2>
                            <p>{{ next_steps }}</p>
                        </div>
                        
                        <div class="footer">
                            <p>Generated by CrewAI Project Documentation Generator</p>
                        </div>
                    </div>
                </body>
                </html>
                """
                template = jinja2.Template(template_str)
            
            # Render HTML
            html_content = template.render(
                title=documentation['title'],
                date=datetime.now().strftime("%Y-%m-%d"),
                executive_summary=documentation['executive_summary'],
                project_scope=documentation['project_scope'],
                stakeholder_analysis=documentation['stakeholder_analysis'],
                functional_requirements=documentation['functional_requirements'],
                technical_requirements=documentation['technical_requirements'],
                timeline=documentation['timeline'],
                budget=documentation['budget'],
                risk_assessment=documentation['risk_assessment'],
                assumptions=documentation['assumptions'],
                next_steps=documentation['next_steps']
            )
            
            # Save HTML file
            async with asyncio.Lock():
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            
            logger.info(f"HTML document generated successfully: {file_path}")
            
            # Create download URL
            download_url = f"{self.base_url}/static/downloads/{filename}"
            
            # Store download info
            await store_download_info({
                "file_id": documentation['file_id'],
                "documentation_id": documentation['documentation_id'],
                "format": "html",
                "file_path": file_path,
                "download_url": download_url,
                "expiry_time": datetime.now().isoformat()
            })
            
            return {
                "success": True,
                "file_path": file_path,
                "download_url": download_url
            }
            
        except Exception as e:
            logger.error(f"Error generating HTML document: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating HTML document: {str(e)}"
            }
    
    async def generate_document(self, documentation_id: str, format: str = "pdf") -> Dict[str, Any]:
        """
        Generate a document in the specified format.
        
        Args:
            documentation_id: ID of the documentation to generate
            format: Document format (pdf, docx, html)
            
        Returns:
            dict: Result of the operation with success status and download_url
        """
        try:
            # Get documentation
            documentation = await get_documentation_by_id(documentation_id)
            if not documentation:
                return {
                    "success": False,
                    "message": "Documentation not found"
                }
            
            # Generate document based on format
            if format == "pdf":
                result = await self.generate_pdf(documentation)
            elif format == "docx":
                result = await self.generate_docx(documentation)
            elif format == "html":
                result = await self.generate_html(documentation)
            else:
                return {
                    "success": False,
                    "message": f"Unsupported format: {format}"
                }
            
            if result["success"]:
                return {
                    "success": True,
                    "documentation_id": documentation_id,
                    "file_id": documentation["file_id"],
                    "download_url": result["download_url"],
                    "format": format
                }
            else:
                return result
                
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating document: {str(e)}"
            }
